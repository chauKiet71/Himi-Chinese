from pathlib import Path
from copy import deepcopy
from zipfile import ZipFile, ZIP_DEFLATED
from collections import Counter
import json, re, hashlib
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

ROOT=Path(__file__).resolve().parents[2]
TMP=Path(__file__).resolve().parent
OUT=ROOT/'outputs'/'noi-dung-bai-hoc-word'
REF=Path('C:/Users/DELL/.codex/plugins/cache/openai-curated-remote/openai-templates/0.1.1/skills/artifact-template-system-design/assets/reference.docx')
D=json.loads((TMP/'content-complete.json').read_text(encoding='utf-8'))
OUT.mkdir(parents=True,exist_ok=True)
parts=[]
def part(filename,title,description):
    b=[]; parts.append(dict(filename=filename,title=title,description=description,blocks=b,lessons=0)); return b
def h(b,t,n=1): b.append(('h',str(t),n))
def p(b,t):
    if t is not None and str(t).strip(): b.append(('p',str(t)))
def rec(b,label,items):
    p(b,label+'\n'+'\n'.join(str(i) for i in items if i is not None and str(i).strip()))
def table(b,headers,rows):
    if rows: b.append(('t',headers,[[str(v) if v is not None else '' for v in row] for row in rows]))
def seq(b,title,items):
    if items:
        h(b,title,3)
        for i,x in enumerate(items,1): p(b,f'{i}. {x}')
def vocab(b,items):
    if not items:return
    h(b,'Từ vựng',3)
    for i,w in enumerate(items,1):
        rec(b,f"{i}. {w.get('hanzi','')}  |  {w.get('pinyin','')}  |  {w.get('meaning','')}",[
            w.get('wordClass'),w.get('example'),w.get('examplePinyin'),w.get('translation'),
            ('Audio: '+w['audioUrl']) if w.get('audioUrl') else None])
        for r in w.get('radicals',[]):p(b,f"Bộ thủ: {r['glyph']} · {r['name']} · {r['strokes']} nét. {r['note']}")
def lines(b,items):
    for x in items: rec(b,f"{x.get('speaker',x.get('role',''))}: {x.get('hanzi','')}",[x.get('pinyin'),x.get('translation')])
def exercises(b,items):
    if not items:return
    h(b,'Bài tập và đáp án',3)
    for i,e in enumerate(items,1):
        rec(b,f"Câu {i}",[e.get('instruction'),e.get('eyebrow'),e.get('prompt'),e.get('chinese'),e.get('pinyin'),e.get('note')])
        if e.get('speakText'):p(b,'Nội dung nghe: '+e['speakText'])
        for j,o in enumerate(e.get('options',[])):p(b,f'{chr(65+j)}. {o}')
        if isinstance(e.get('correctOption'),int):
            n=e['correctOption'];p(b,f"Đáp án: {chr(65+n)}. {e['options'][n]}")
        elif e.get('answer') is not None:p(b,'Đáp án: '+str(e['answer']))
        else:p(b,'Đáp án chưa có trong dữ liệu bài học.')
        p(b,e.get('explanation'))
        if e.get('audioUrl'):p(b,'Audio: '+e['audioUrl'])
def lesson(b,title,summary='',n=2):
    h(b,title,n);p(b,summary);parts[-1]['lessons']+=1

for idx,c in enumerate(D['courses'],1):
    b=part(f'{idx:02d}_{c["slug"]}.docx',c['title'],c['description'])
    p(b,c['chineseTitle']); count=0
    for mod in c['modules']:
        h(b,mod['title']);p(b,mod['description'])
        for l in c['lessons']:
            if l['moduleSlug']!=mod['slug']:continue
            count+=1;lesson(b,f"Bài {count}  {l['title']}",l['summary'])
            p(b,f"Tình huống: {l['situation']}\nThời lượng: {l['estimatedMinutes']} phút")
            vocab(b,l['vocabulary']);content=l['content']
            for key,title in [('dialogue','Hội thoại'),('phrases','Câu mẫu')]:
                if content.get(key):h(b,title,3);lines(b,content[key])
            if content['notes']:
                h(b,'Cách dùng và lưu ý',3)
                for note in content['notes']:rec(b,note['title'],[note['pattern'],note['explanation']])
            if content.get('challenge'):
                ch=content['challenge'];h(b,ch['title'],3);p(b,ch['description']);p(b,f"Điểm đạt: {ch['passScore']}");exercises(b,ch['questions'])
    assert count==len(c['lessons'])

raw5=json.loads((ROOT/'content/hsk5-workbook-1-json/exercises.json').read_text(encoding='utf-8'))['exercises']
blocks5=[]
for name in ['listening','reading','writing']:
    blocks5.extend(json.loads((ROOT/f'content/hsk5-workbook-1-json/blocks/{name}-blocks.json').read_text(encoding='utf-8'))['blocks'])
for idx,level in enumerate(D['hsk'],9):
    b=part(f'{idx:02d}_{level["id"]}.docx',level['label'],level['description'])
    for topic in level['topics']:
        h(b,topic['title'])
        for l in topic['lessons']:
            c=l.get('content');lesson(b,f"Bài {l['lessonNumber']}  {l['title']}",c['summary'] if c else '')
            if not c:p(b,'Bài học hiện có tên trong lộ trình nhưng chưa có nội dung chi tiết.');continue
            p(b,f"Thời lượng: {c['minutes']} phút");p(b,c.get('greeting'))
            vocab(b,c['vocabulary'])
            if c['grammar']:
                h(b,'Ngữ pháp',3)
                for g in c['grammar']:
                    rec(b,g['title'],[g['formula'],g['explanation']])
                    for ex in g['examples']:rec(b,ex['hanzi'],[ex['pinyin'],ex['translation']])
            if c['dialogues']:
                h(b,'Hội thoại và bài đọc',3)
                for dialog in c['dialogues']:rec(b,dialog['title'],[dialog['setting']]);lines(b,dialog['turns'])
            seq(b,'Phát âm',c['pronunciationTopics'])
            exercises(b,c['exercises'])
            if c['writingCharacters']:
                h(b,'Chữ Hán luyện viết',3)
                table(b,['Chữ','Từ gốc','Pinyin','Nghĩa'],[[w['hanzi'],w['word'],w['pinyin'],w['meaning']] for w in c['writingCharacters']])
            if c['sourceId'].startswith('hsk5w1-'):
                h(b,'Nội dung sách bài tập gốc',3)
                for block in blocks5:
                    if block['lessonId']!=c['sourceId']:continue
                    p(b,block.get('contentOcr') or '\n'.join(block.get('contentOcrLines',[])))
                    if block.get('audioTrackReferences'):p(b,'Tham chiếu audio: '+', '.join(block['audioTrackReferences']))
                for e in raw5:
                    if e['lessonId']!=c['sourceId']:continue
                    p(b,f"Câu {e['numberInSource']} · {dict(listening='Nghe',reading='Đọc',writing='Viết')[e['sectionType']]} · Phần {e['part']}")
                    p(b,e['promptOcr'])
                    for opt in e['optionsOcr']:p(b,opt['label']+'. '+opt['textOcr'])

b=part('16_hsk-2-sach-bai-tap.docx','HSK 2 Sách bài tập','Nội dung nghe hiểu, đọc hiểu, phát âm và chữ Hán theo từng bài trong sách bài tập HSK 2.')
for l in D['workbook2']:
    lesson(b,f"Bài {l['lessonNumber']}  {l['title']}",l['summary'],1);p(b,l['titleZh']);p(b,l['pinyin'])
    for section in l['sections']:
        h(b,section['label'],2)
        for g in section['groups']:
            h(b,g['exerciseLabel']+(f" phần {g['part']}" if g['part'] else ''),3)
            p(b,g['instructionZh']);p(b,g['instructionVi'])
            if g['requiresAudio'] and not g['audioAvailable']:p(b,'Bài yêu cầu audio nhưng nguồn hiện chưa cung cấp audio.')
            if g['requiresVisual'] and not g['visualAvailable']:p(b,'Bài yêu cầu hình ảnh nhưng nguồn hiện chưa cung cấp hình ảnh.')
            for opt in g['optionBank']:rec(b,opt['label']+'. '+opt['text'],[opt['pinyin']])
            for item in g['items']:
                rec(b,f"Câu {item['questionNumber']}",item['chineseSegments']+item['pinyinSegments'])
                for opt in item['options']:rec(b,opt['label']+'. '+opt['text'],[opt['pinyin']])
                if not item['answerAvailable']:p(b,'Chưa có đáp án trong dữ liệu bài học.')
            if g['sourceLines']:
                p(b,'Văn bản gốc được nhận dạng từ trang sách')
                for line in g['sourceLines']:p(b,line['text'])

b=part('17_luyen-tap-tinh-huong.docx','Luyện tập theo tình huống','Các tình huống giao tiếp thực tế và bài tập kèm đáp án theo từng ngành nghề.')
for ind in D['practice']:
    h(b,ind['label']);p(b,ind['description'])
    for l in ind['lessons']:
        lesson(b,l['title'],l['brief']);p(b,l['context']);p(b,f"Trình độ: {l['level']} · Thời lượng: {l['durationMinutes']} phút")
        rec(b,l['sentenceZh'],[l['pinyin'],l['translation']]);seq(b,'Trọng tâm',l['focus']);exercises(b,l['exercises'])
        h(b,'Bài nghe đúng sai và câu hỏi nghĩa',3)
        for i,e in enumerate(l['exercises'],1):
            s=e['listeningStatement'];rec(b,f"Câu {i}",[s['text'],'Kết quả: '+('Đúng' if s['isCorrect'] else 'Sai'),'Cách nói đúng: '+s['correctText'],s['explanation']])
            q=e['meaningQuestion']
            for j,o in enumerate(q['options']):p(b,f'{chr(65+j)}. {o}')
            p(b,'Đáp án nghĩa: '+chr(65+q['correctOption'])+'. '+q['options'][q['correctOption']])

b=part('18_luyen-nghe.docx','Luyện nghe','Từ vựng, pinyin, nghĩa và câu ví dụ của từng bài luyện nghe theo cấp độ.')
for level in D['listening']:
    h(b,level['label']+' '+level['title']);p(b,level['description'])
    for l in level['lessons']:
        lesson(b,f"Bài {l['order']}  {l['title']}",l['description']);vocab(b,l['words'])

b=part('19_luyen-viet.docx','Luyện viết chữ Hán','Danh sách chữ Hán, pinyin và nghĩa để luyện viết theo từng cấp độ và bài học.')
for level in D['writing']:
    h(b,level['label']);p(b,level['description'])
    for l in level['lessons']:
        lesson(b,f"{l['sourceLabel']} Bài {l['lessonNumber']}  {l['title']}",l['summary']);p(b,l['duration']);seq(b,'Mục tiêu',l['outcomes'])
        table(b,['Chữ Hán','Pinyin','Nghĩa'],[[w['hanzi'],w['pinyin'],w['meaning']] for w in l['characters']])

b=part('20_hoc-qua-video.docx','Học qua video','Nội dung bài học, mục tiêu, gợi ý luyện tập và lời thoại có sẵn trong thư viện video.')
for v in D['videos']:
    lesson(b,v['title'],v['summary'],1);p(b,v.get('originalTitle'));p(b,v['description']);p(b,f"Chủ đề: {v['category']} · Trình độ: {v['level']}")
    p(b,v.get('durationLabel'))
    if v.get('authorName'):p(b,'Tác giả: '+v['authorName'])
    if v.get('authorUrl'):p(b,v['authorUrl'])
    if v.get('youtubeId'):p(b,'Video: https://www.youtube.com/watch?v='+v['youtubeId'])
    elif v.get('videoUrl'):p(b,'Video: '+v['videoUrl'])
    seq(b,'Mục tiêu học tập',v['learningGoals']);seq(b,'Gợi ý luyện tập',v['practicePrompts'])
    if v.get('transcript'):
        h(b,'Lời thoại',2)
        for x in v['transcript']:
            t=x['startMs']//1000;p(b,f"{t//60:02d}:{t%60:02d} · {x['role']}");rec(b,x['hanzi'],[x['pinyin'],x['translation'],'Từ khóa: '+x['keyword']])
    else:p(b,'Video này chưa có lời thoại trong dữ liệu bài học.')

b=part('21_tro-choi-tu-vung.docx','Trò chơi từ vựng','Toàn bộ từ vựng và câu ví dụ trong ngân hàng từ của trò chơi.')
vocab(b,D['games']);parts[-1]['lessons']=len(D['games'])

def newdoc(title,description):
    d=Document(REF)
    for el in list(d._element.body):
        if el.tag!=qn('w:sectPr'):d._element.body.remove(el)
    d.add_paragraph('Himi Chinese',style='Subtitle')
    d.add_paragraph(title,style='Title')
    d.add_paragraph(description,style='normal')
    return d
def fmt(p):
    for r in p.runs:
        fonts=r._element.get_or_add_rPr().find(qn('w:rFonts'))
        if fonts is None:fonts=OxmlElement('w:rFonts');r._element.get_or_add_rPr().insert(0,fonts)
        fonts.set(qn('w:eastAsia'),'SimSun')
    return p
def addblocks(d,blocks):
    for block in blocks:
        if block[0]=='h':
            q=d.add_paragraph(block[1],style=f'Heading {block[2]}');q.paragraph_format.keep_with_next=True
            fmt(q)
        elif block[0]=='p':fmt(d.add_paragraph(block[1],style='normal'))
        else:
            headers,rows=block[1:];t=d.add_table(rows=1,cols=len(headers));t.autofit=False
            widths=([.75,1.0,1.35,4.0] if len(headers)==4 else [1.0,1.55,4.55])
            for col,width in zip(t.columns,widths):col.width=Inches(width)
            props=t._tbl.tblPr
            borders=OxmlElement('w:tblBorders')
            for edge in ['top','left','bottom','right','insideH','insideV']:
                e=OxmlElement('w:'+edge);e.set(qn('w:val'),'single');e.set(qn('w:sz'),'4');e.set(qn('w:color'),'D9D9D9');borders.append(e)
            props.append(borders)
            margins=OxmlElement('w:tblCellMar')
            for edge in ['top','left','bottom','right']:
                e=OxmlElement('w:'+edge);e.set(qn('w:w'),'80');e.set(qn('w:type'),'dxa');margins.append(e)
            props.append(margins)
            repeat=OxmlElement('w:tblHeader');t.rows[0]._tr.get_or_add_trPr().append(repeat)
            for j,name in enumerate(headers):
                cell=t.rows[0].cells[j];cell.text=name;shade=OxmlElement('w:shd');shade.set(qn('w:fill'),'CFE2F3');cell._tc.get_or_add_tcPr().append(shade)
                for p0 in cell.paragraphs:
                    for r in p0.runs:r.bold=True
                    fmt(p0)
            for row in rows:
                cells=t.add_row().cells
                for j,val in enumerate(row):cells[j].text=val;cells[j].width=Inches(widths[j]);fmt(cells[j].paragraphs[0])
            d.add_paragraph('',style='normal')
    return d
def save_preserving(d,path):
    # Preserve every original package member except the replaced document body.
    body=etree.tostring(d._element,xml_declaration=True,encoding='UTF-8',standalone=True)
    with ZipFile(REF) as src, ZipFile(path,'w',ZIP_DEFLATED) as dst:
        for info in src.infolist():dst.writestr(info,body if info.filename=='word/document.xml' else src.read(info.filename))
    return body

refsha=hashlib.sha256(REF.read_bytes()).hexdigest()
contract=f'''# Artifact template contract
Reference: {REF}\nSHA256: {refsha}
One portrait section 8.5 by 11 inches. Margins 0.7 top left right and 0.620139 bottom. Preserve section properties exactly.
Preserve template styles, theme, numbering, package relationships, headers and footers byte for byte. Reuse Title, Subtitle, normal, Heading 1 through 3. Title 22 pt and Heading 1 13.5 pt from retained styles. Preserve inherited typography and colors. Add only SimSun East Asian fallback to generated runs so Chinese is supported.
Replace all template body prose and sample diagrams with curriculum content. Repeated modules and lessons expand the body slot. Optional proposal-specific metadata and architecture samples are removed since they do not apply to lesson export. Tables repeat source pale-blue color and use light-gray borders, flexible row heights and repeated headers.
No fields introduced. Original section and page furniture retained. Source package and all members except word/document.xml must remain unchanged. User content may extend to any number of pages.
Visual evidence: assets/preview.png inspected. Canonical renderer attempt failed because soffice.exe is unavailable. Native Word is also unavailable. Page count and full visual QA remain unverified; disclose limitation on delivery.
'''
(TMP/'artifact.md').write_text(contract,encoding='utf-8')
manifest=[]
for item in parts:
    d=newdoc(item['title'],item['description']);addblocks(d,item['blocks']);xml=save_preserving(d,OUT/item['filename'])
    manifest.append({k:item[k] for k in ['filename','title','lessons']}|{'paragraphs':len(d.paragraphs),'sha256':hashlib.sha256((OUT/item['filename']).read_bytes()).hexdigest()})
    print('Created',item['filename'],flush=True)

master=newdoc('Toàn bộ nội dung bài học tiếng Trung','Tổng hợp các khóa học giao tiếp và ngành nghề, HSK, sách bài tập, luyện tập tình huống, luyện nghe, luyện viết, học qua video và trò chơi từ vựng.')
master.add_paragraph('Danh mục các phần',style='Heading 1')
for i,item in enumerate(parts,1):master.add_paragraph(f"{i}. {item['title']}",style='normal')
for i,item in enumerate(parts,1):
    master.add_page_break();master.add_paragraph(f"Phần {i}  {item['title']}",style='Title');master.add_paragraph(item['description'],style='normal');addblocks(master,item['blocks'])
save_preserving(master,OUT/'00_toan-bo-noi-dung-bai-hoc.docx')

# Verify every emitted content string survives OOXML round trip; preserve template package.
checks=[]
for item in parts:
    with ZipFile(OUT/item['filename']) as z,ZipFile(REF) as ref:
        assert z.testzip() is None
        root=etree.fromstring(z.read('word/document.xml'))
        texts='\n'.join(root.xpath('//w:t/text()',namespaces={'w':qn('w:t').split('}')[0][1:]}))
        expected=[]
        for block in item['blocks']:
            if block[0]=='t':expected.extend(block[1]);expected.extend(v for row in block[2] for v in row)
            else:expected.append(block[1])
        for text in expected:
            for line in text.replace('\r','').split('\n'):
                for segment in line.split('\t'):assert not segment or segment in texts,(item['filename'],segment[:100])
        for name in ref.namelist():
            if name!='word/document.xml':assert z.read(name)==ref.read(name),(name,'changed')
        assert not re.search(r'\[(Summarize|Describe|Open question|State the|Name\(s\))',texts)
        checks.append({'file':item['filename'],'verified_strings':len(expected),'package_preserved':True})
with ZipFile(OUT/'00_toan-bo-noi-dung-bai-hoc.docx') as z:
    masterroot=etree.fromstring(z.read('word/document.xml'));assert z.testzip() is None
    mastertexts='\n'.join(masterroot.xpath('//w:t/text()',namespaces={'w':qn('w:t').split('}')[0][1:]}))
    for item in parts:
        for block in item['blocks']:
            values=block[1]+[v for row in block[2] for v in row] if block[0]=='t' else [block[1]]
            for value in values:
                for line in value.replace('\r','').split('\n'):
                    for segment in line.split('\t'):assert not segment or segment in mastertexts,segment[:100]
assert hashlib.sha256(REF.read_bytes()).hexdigest()==refsha
(TMP/'manifest.json').write_text(json.dumps(manifest,ensure_ascii=False,indent=2),encoding='utf-8')
(TMP/'verification.json').write_text(json.dumps(checks,ensure_ascii=False,indent=2),encoding='utf-8')
readme='BỘ NỘI DUNG BÀI HỌC HIMI CHINESE\nNgày xuất 06/09/2026\n\n00_toan-bo-noi-dung-bai-hoc.docx là bản tổng hợp.\nCác file 01 đến 21 là bản riêng cho từng phần.\n\n'
readme+='\n'.join(f"{m['filename']} — {m['title']} ({m['lessons']} bài hoặc mục)" for m in manifest)
readme+='\n\nPhạm vi: dữ liệu bài học trong dự án cục bộ tại thời điểm xuất; chưa đối chiếu cơ sở dữ liệu vận hành.\nHSK 7–9 hiện chỉ có danh mục bài. Một số nội dung workbook thiếu audio, hình ảnh hoặc đáp án và còn ở dạng OCR. Tài liệu giữ nguyên các phần hiện có.\nAudio và video được lưu dưới dạng địa chỉ khi có, không nhúng tệp đa phương tiện. Đường dẫn bắt đầu bằng / là tài nguyên của ứng dụng.\nNội dung luyện viết được xuất thành chữ và bảng từ; hoạt ảnh thứ tự nét và chấm điểm tương tác sử dụng trên ứng dụng.\nĐã kiểm tra nội dung và cấu trúc DOCX tự động. Chưa kiểm tra được bố cục từng trang vì môi trường không có bộ dựng trang Word/LibreOffice.\n'
(OUT/'HUONG-DAN.txt').write_text(readme,encoding='utf-8-sig')
with ZipFile(ROOT/'outputs/Toan-bo-bai-hoc-Word.zip','w',ZIP_DEFLATED) as z:
    for path in sorted(OUT.iterdir()):z.write(path,path.name)
print(json.dumps({'docx_files':len(parts)+1,'verified_strings':sum(x['verified_strings'] for x in checks),'zip':str(ROOT/'outputs/Toan-bo-bai-hoc-Word.zip')},ensure_ascii=False))
